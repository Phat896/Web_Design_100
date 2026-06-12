from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "artifacts" / "report-assets"
TEMPLATE_MEDIA = ROOT / "artifacts" / "template-media"
OUTPUT = ROOT / "artifacts" / "Bao_cao_StellarMind_Nguyen_Tan_Phat_25AS037_CAP_NHAT.docx"

TITLE = (
    "XÂY DỰNG WEBSITE STELLARMIND – HỆ THỐNG KHÁM PHÁ HỆ MẶT TRỜI "
    "TƯƠNG TÁC KẾT HỢP HỖ TRỢ CẢM XÚC"
)
SHORT_TITLE = "Báo cáo dự án StellarMind"
STUDENT = "Nguyễn Tấn Phát"
STUDENT_ID = "25AS037"
STUDENT_CLASS = "25AS"
SUPERVISOR = "ThS. Nguyễn Ngọc Huyền Trân"

NAVY = "17365D"
BLUE = "1D6FA5"
LIGHT_BLUE = "EAF3F8"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D8E0E6"
TEXT_GRAY = "4B5563"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches):
    width = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_fixed_layout(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_inches in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width_inches * 1440)))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Times New Roman", size=13, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_keep(paragraph, keep_with_next=False, keep_together=False):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_with_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_together:
        p_pr.append(OxmlElement("w:keepLines"))


def add_field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def set_update_fields(document):
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def set_page_number_format(section, fmt="decimal", start=None):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    add_field(p, "PAGE")
    for run in p.runs:
        set_run_font(run, size=11)


def add_header(section, text):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run, size=9, italic=True, color=TEXT_GRAY)
    p.paragraph_format.space_after = Pt(0)


def set_section_geometry(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.9)


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.27)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Heading 1", 16, NAVY),
        ("Heading 2", 14, NAVY),
        ("Heading 3", 13, BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Heading 1"].font.all_caps = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(13)
        style.paragraph_format.left_indent = Cm(1.27)
        style.paragraph_format.first_line_indent = Cm(-0.63)
        style.paragraph_format.line_spacing = 1.2
        style.paragraph_format.space_after = Pt(3)

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(11)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(TEXT_GRAY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    if "Figure Caption" not in styles:
        styles.add_style("Figure Caption", 1)
    figure_caption = styles["Figure Caption"]
    figure_caption.base_style = caption

    if "Table Caption" not in styles:
        styles.add_style("Table Caption", 1)
    table_caption = styles["Table Caption"]
    table_caption.base_style = caption

    if "Code" not in styles:
        styles.add_style("Code", 1)
    code_style = styles["Code"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(9)
    code_style.paragraph_format.first_line_indent = Cm(0)
    code_style.paragraph_format.left_indent = Cm(0.5)
    code_style.paragraph_format.right_indent = Cm(0.5)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.line_spacing = 1.0
    p_pr = code_style._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)


def add_body(document, text, bold_prefix=None):
    p = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)


def add_heading(document, text, level):
    p = document.add_paragraph(text, style=f"Heading {level}")
    for run in p.runs:
        set_run_font(
            run,
            size={1: 16, 2: 14, 3: 13}[level],
            bold=True,
            color=NAVY if level < 3 else BLUE,
        )
    return p


def add_chapter(document, title):
    document.add_page_break()
    return add_heading(document, title, 1)


def add_callout(document, title, text):
    table = document.add_table(rows=1, cols=1)
    set_table_fixed_layout(table, [6.2])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2)
    title_run = p.add_run(title + ": ")
    set_run_font(title_run, size=12, bold=True, color=NAVY)
    text_run = p.add_run(text)
    set_run_font(text_run, size=12)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(document, caption, headers, rows, widths):
    cap = document.add_paragraph(caption, style="Table Caption")
    set_paragraph_keep(cap, keep_with_next=True)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_fixed_layout(table, widths)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, text in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(text))
        set_run_font(run, size=10.5, bold=True, color=WHITE)

    for row_index, values in enumerate(rows):
        row_cells = table.add_row().cells
        if row_index % 2:
            for cell in row_cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for index, value in enumerate(values):
            p = row_cells[index].paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.1
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_run_font(run, size=10.5)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_figure(document, path, caption, width=6.1):
    intro_guard = document.paragraphs[-1] if document.paragraphs else None
    if intro_guard:
        set_paragraph_keep(intro_guard, keep_with_next=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    picture = run.add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    set_paragraph_keep(p, keep_with_next=True)
    cap = document.add_paragraph(caption, style="Figure Caption")
    return cap


def add_cover(document, secondary=False):
    section = document.sections[-1]
    set_section_geometry(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].clear()
    section.footer.paragraphs[0].clear()

    logo_path = TEMPLATE_MEDIA / "image1.png"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    logo = p.add_run().add_picture(str(logo_path), width=Inches(6.2))
    logo._inline.docPr.set("descr", "Biểu trưng Trường Đại học Công nghệ Thông tin và Truyền thông Việt - Hàn")
    logo._inline.docPr.set("title", "Biểu trưng VKU")

    for text, size, bold, spacing in [
        ("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN\nVÀ TRUYỀN THÔNG VIỆT – HÀN", 14, True, 2),
        ("KHOA KỸ THUẬT MÁY TÍNH", 14, True, 8),
        ("DỰ ÁN", 18, True, 18),
        (TITLE, 18, True, 18),
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(spacing)
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=NAVY if "STELLARMIND" in text else None)

    for _ in range(2 if secondary else 3):
        document.add_paragraph()

    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed_layout(table, [2.25, 3.75])
    details = [
        ("Sinh viên thực hiện:", STUDENT),
        ("Lớp:", STUDENT_CLASS),
        ("Mã sinh viên:", STUDENT_ID),
        ("Giảng viên hướng dẫn:", SUPERVISOR),
    ]
    for row, (label, value) in zip(table.rows, details):
        for cell in row.cells:
            cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.first_line_indent = Cm(0)
        r1 = p1.add_run(label)
        set_run_font(r1, size=13, bold=True)
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.first_line_indent = Cm(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=13)

    for _ in range(3):
        document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("Đà Nẵng, tháng 06 năm 2026")
    set_run_font(run, size=13, italic=True)


def add_front_matter(document):
    document.add_page_break()
    add_heading(document, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", 1)
    for _ in range(20):
        p = document.add_paragraph("........................................................................................................................")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(7)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=12, color=TEXT_GRAY)

    front_section = document.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(front_section)
    set_page_number_format(front_section, fmt="lowerRoman", start=1)
    add_page_number(front_section)
    add_header(front_section, SHORT_TITLE)

    add_heading(document, "LỜI CẢM ƠN", 1)
    for text in [
        (
            "Trước hết, em xin chân thành cảm ơn quý thầy cô Trường Đại học Công nghệ Thông tin "
            "và Truyền thông Việt – Hàn, đặc biệt là các thầy cô Khoa Kỹ thuật Máy tính, đã cung cấp "
            "nền tảng kiến thức về lập trình web, cơ sở dữ liệu và phân tích thiết kế hệ thống."
        ),
        (
            f"Em xin gửi lời cảm ơn sâu sắc đến {SUPERVISOR}, người đã hướng dẫn, góp ý và hỗ trợ "
            "em trong quá trình lựa chọn hướng tiếp cận, hoàn thiện sản phẩm và xây dựng báo cáo dự án."
        ),
        (
            "Em cũng cảm ơn bạn bè và những người đã dành thời gian trải nghiệm giao diện, chia sẻ "
            "nhận xét về khả năng sử dụng và động viên em trong quá trình thực hiện dự án StellarMind."
        ),
        (
            "Do giới hạn về thời gian, kinh nghiệm phát triển và điều kiện kiểm thử, sản phẩm vẫn còn "
            "một số hạn chế. Em kính mong nhận được các ý kiến đóng góp để tiếp tục hoàn thiện hệ thống "
            "theo hướng ổn định, an toàn và có giá trị giáo dục cao hơn."
        ),
    ]:
        add_body(document, text)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run("Sinh viên thực hiện\n").bold = True
    p.add_run(STUDENT)
    for run in p.runs:
        set_run_font(run, size=13)

    document.add_page_break()
    add_heading(document, "MỤC LỤC", 1)
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "Mục lục sẽ được cập nhật khi mở tài liệu.")

    document.add_page_break()
    add_heading(document, "DANH MỤC HÌNH ẢNH", 1)
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    add_field(p, 'TOC \\h \\z \\t "Figure Caption,1"', "Danh mục hình sẽ được cập nhật khi mở tài liệu.")

    document.add_page_break()
    add_heading(document, "DANH MỤC BẢNG BIỂU", 1)
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    add_field(p, 'TOC \\h \\z \\t "Table Caption,1"', "Danh mục bảng sẽ được cập nhật khi mở tài liệu.")

    document.add_page_break()
    add_heading(document, "DANH MỤC TỪ VIẾT TẮT", 1)
    add_table(
        document,
        "Bảng 0.1. Danh mục từ viết tắt",
        ["Từ viết tắt", "Tên đầy đủ", "Ý nghĩa sử dụng trong báo cáo"],
        [
            ("API", "Application Programming Interface", "Giao diện lập trình ứng dụng"),
            ("CSS", "Cascading Style Sheets", "Ngôn ngữ định kiểu giao diện web"),
            ("DOM", "Document Object Model", "Mô hình đối tượng tài liệu HTML"),
            ("HTML", "HyperText Markup Language", "Ngôn ngữ đánh dấu siêu văn bản"),
            ("HTTP", "HyperText Transfer Protocol", "Giao thức trao đổi dữ liệu web"),
            ("JSON", "JavaScript Object Notation", "Định dạng dữ liệu trao đổi giữa frontend và backend"),
            ("JWT", "JSON Web Token", "Mã thông báo dùng cho xác thực"),
            ("REST", "Representational State Transfer", "Phong cách thiết kế API"),
            ("UI", "User Interface", "Giao diện người dùng"),
            ("UX", "User Experience", "Trải nghiệm người dùng"),
        ],
        [1.05, 2.2, 2.95],
    )


def add_introduction(document):
    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(body_section)
    set_page_number_format(body_section, fmt="decimal", start=1)
    add_page_number(body_section)
    add_header(body_section, SHORT_TITLE)

    add_heading(document, "MỞ ĐẦU", 1)
    add_heading(document, "1. Lý do chọn đề tài", 2)
    for text in [
        (
            "Các website giáo dục thiên văn thường tập trung vào việc trình bày dữ liệu, hình ảnh hoặc "
            "mô phỏng quỹ đạo. Trong khi đó, các nền tảng chia sẻ cảm xúc thường sử dụng giao diện quen "
            "thuộc của mạng xã hội. Đề tài StellarMind được lựa chọn nhằm thử nghiệm một hướng kết hợp "
            "hai nhóm nội dung này trong cùng một trải nghiệm web tương tác."
        ),
        (
            "Hệ Mặt Trời cung cấp một hệ thống hình tượng trực quan, có khả năng tổ chức nội dung theo "
            "tám hành tinh với màu sắc và đặc điểm riêng. Dự án sử dụng mỗi hành tinh như một không gian "
            "chủ đề để người dùng vừa tiếp cận kiến thức thiên văn cơ bản, vừa đọc nội dung tham khảo "
            "liên quan đến một trạng thái cảm xúc."
        ),
        (
            "Về mặt kỹ thuật, đề tài tạo điều kiện vận dụng đồng thời các kiến thức về Canvas, thiết kế "
            "giao diện đáp ứng, lập trình JavaScript phía trình duyệt, xây dựng REST API, xác thực người "
            "dùng và lưu trữ dữ liệu NoSQL. Đây là phạm vi phù hợp với một đồ án cơ sở có cả frontend "
            "và backend nhưng vẫn có thể triển khai theo từng module độc lập."
        ),
    ]:
        add_body(document, text)

    add_heading(document, "2. Mục tiêu đề tài", 2)
    add_body(
        document,
        "Mục tiêu tổng quát của đề tài là xây dựng một website tương tác có khả năng mô phỏng Hệ Mặt Trời, "
        "tổ chức nội dung theo tám hành tinh và cung cấp các công cụ chia sẻ, tương tác cộng đồng ở mức cơ bản.",
    )
    add_bullets(
        document,
        [
            "Xây dựng mô phỏng quỹ đạo và giao diện khám phá hành tinh bằng HTML5 Canvas.",
            "Thiết kế trang nội dung riêng cho tám hành tinh, có thông tin thiên văn và nội dung hỗ trợ cảm xúc.",
            "Cung cấp bài tập thở, âm thanh và các hoạt động tương tác phía trình duyệt.",
            "Xây dựng chức năng đăng ký bằng email, đăng nhập, đăng bài, trả lời, thích và lọc bài viết.",
            "Tích hợp chatbot Gemini để cung cấp phản hồi tham khảo theo ngữ cảnh của dự án.",
            "Xây dựng luồng gửi liên hệ qua backend và dịch vụ SMTP khi máy chủ được cấu hình.",
            "Tổ chức mã nguồn theo hướng tách frontend, backend, model và route để thuận tiện bảo trì.",
        ],
    )

    add_heading(document, "3. Đối tượng và phạm vi nghiên cứu", 2)
    add_body(
        document,
        "Đối tượng nghiên cứu gồm kỹ thuật xây dựng giao diện web tương tác, mô phỏng đồ họa hai chiều bằng "
        "Canvas, kiến trúc REST API, cơ chế xác thực JWT, lưu trữ MongoDB và tích hợp dịch vụ sinh nội dung.",
    )
    add_body(
        document,
        "Phạm vi sản phẩm bao gồm trang chủ, tám trang hành tinh, trang cộng đồng, cửa sổ chatbot và trang "
        "liên hệ. Hệ thống chưa bao gồm phân quyền quản trị, kiểm duyệt nội dung, phân tích tâm lý chuyên môn "
        "hoặc chẩn đoán sức khỏe. Nội dung cảm xúc chỉ mang tính tham khảo và hỗ trợ trải nghiệm.",
    )

    add_heading(document, "4. Phương pháp thực hiện", 2)
    add_numbered(
        document,
        [
            "Khảo sát yêu cầu và xác định phạm vi chức năng phù hợp với thời gian thực hiện.",
            "Thiết kế giao diện, luồng tương tác và cấu trúc dữ liệu trước khi triển khai.",
            "Phát triển từng module frontend, sau đó tích hợp backend và cơ sở dữ liệu.",
            "Kiểm tra cú pháp, mở các trang trên máy chủ cục bộ và ghi nhận các chức năng cần kiểm thử tiếp.",
            "Đối chiếu mã nguồn với báo cáo để loại bỏ nội dung không có bằng chứng hoặc chưa hoàn thiện.",
        ],
    )

    add_heading(document, "5. Nội dung và kế hoạch thực hiện", 2)
    add_table(
        document,
        "Bảng 0.2. Kế hoạch thực hiện đề tài",
        ["Giai đoạn", "Nội dung chính", "Sản phẩm đầu ra"],
        [
            ("1. Khảo sát", "Xác định chủ đề, người dùng và chức năng cốt lõi", "Danh sách yêu cầu"),
            ("2. Thiết kế", "Thiết kế UI, kiến trúc, model và route", "Sơ đồ và cấu trúc dữ liệu"),
            ("3. Frontend", "Mô phỏng Canvas, trang hành tinh, cộng đồng, liên hệ", "Giao diện có thể chạy cục bộ"),
            ("4. Backend", "Xác thực, bình luận, chatbot, liên hệ SMTP và kết nối MongoDB", "REST API Express"),
            ("5. Hoàn thiện", "Kiểm tra, sửa lỗi, chụp giao diện và viết báo cáo", "Sản phẩm và tài liệu đồ án"),
        ],
        [1.2, 2.9, 2.1],
    )

    add_heading(document, "6. Bố cục báo cáo", 2)
    add_body(
        document,
        "Ngoài phần mở đầu, kết luận, tài liệu tham khảo và phụ lục, báo cáo được tổ chức thành ba chương. "
        "Chương 1 trình bày tổng quan và công nghệ. Chương 2 tập trung vào phân tích, thiết kế, dữ liệu và "
        "giao diện. Chương 3 mô tả quá trình triển khai, REST API, kiểm thử, kết quả và hạn chế.",
    )


def add_chapter_one(document):
    add_chapter(document, "CHƯƠNG 1: TỔNG QUAN VỀ DỰ ÁN STELLARMIND")

    add_heading(document, "1.1. Giới thiệu đề tài", 2)
    add_body(
        document,
        "StellarMind là website khám phá Hệ Mặt Trời tương tác được xây dựng theo phong cách khoa học viễn "
        "tưởng. Trang chủ thể hiện Mặt Trời, tám hành tinh, quỹ đạo, vành đai tiểu hành tinh và các lớp nền "
        "không gian bằng Canvas. Người dùng có thể chọn hành tinh từ thanh bên hoặc trực tiếp trên mô phỏng "
        "để xem thông tin và chuyển sang trang nội dung chi tiết.",
    )
    add_body(
        document,
        "Bên cạnh dữ liệu thiên văn, mỗi hành tinh được liên kết với một chủ đề cảm xúc. Các trang hành tinh "
        "sử dụng nội dung, màu sắc, hình ảnh, âm thanh và hoạt động tương tác khác nhau. Cách tổ chức này tạo "
        "ra hành trình khám phá có tính kể chuyện thay vì chỉ hiển thị thông tin dưới dạng danh sách.",
    )

    add_heading(document, "1.2. Bài toán và nhu cầu thực tế", 2)
    add_body(
        document,
        "Bài toán đặt ra là xây dựng một sản phẩm web vừa trực quan, vừa có nhiều hình thức tương tác nhưng "
        "không phụ thuộc vào một framework frontend lớn. Hệ thống cần duy trì tốc độ phản hồi tốt, hoạt động "
        "trên trình duyệt phổ biến và cho phép bổ sung nội dung hành tinh mà không thay đổi kiến trúc tổng thể.",
    )
    add_body(
        document,
        "Một nhu cầu khác là tạo không gian để người dùng ghi lại cảm xúc và nhận phản hồi từ cộng đồng. "
        "Do nội dung có tính cá nhân, hệ thống hỗ trợ cả chế độ hiển thị tên tài khoản và chế độ ẩn danh. "
        "Chatbot được bổ sung như một kênh tương tác tham khảo, không thay thế chuyên gia tâm lý hoặc dịch vụ y tế.",
    )

    add_heading(document, "1.3. Mục tiêu của hệ thống", 2)
    add_bullets(
        document,
        [
            "Trình bày Hệ Mặt Trời bằng mô phỏng trực quan và có thể điều khiển.",
            "Cung cấp nội dung nhất quán cho tám hành tinh và tám chủ đề cảm xúc.",
            "Hỗ trợ người dùng thực hiện các hoạt động thư giãn đơn giản trên trình duyệt.",
            "Cho phép tạo tài khoản và tham gia cộng đồng chia sẻ.",
            "Cung cấp API rõ ràng cho xác thực, bình luận và chatbot.",
            "Bảo đảm báo cáo phản ánh đúng trạng thái hiện tại của mã nguồn.",
        ],
    )

    add_heading(document, "1.4. Đối tượng sử dụng", 2)
    add_body(
        document,
        "Hệ thống hướng tới người dùng phổ thông có nhu cầu khám phá thiên văn qua giao diện tương tác, sinh "
        "viên muốn tham khảo cách kết hợp Canvas với REST API, và người dùng muốn ghi lại trạng thái cảm xúc "
        "trong một không gian trực quan nhẹ nhàng.",
    )
    add_callout(
        document,
        "Lưu ý phạm vi",
        "StellarMind không thu thập dữ liệu lâm sàng, không chẩn đoán bệnh và không đưa ra phác đồ điều trị. "
        "Các nội dung về cảm xúc, bài tập thở và phản hồi chatbot chỉ mang tính tham khảo.",
    )

    add_heading(document, "1.5. Phạm vi hệ thống", 2)
    add_body(
        document,
        "Frontend gồm các trang HTML độc lập, stylesheet dùng chung, module JavaScript và tài nguyên hình ảnh. "
        "Tám trang hành tinh hiện dùng chung bộ dữ liệu planet-data.js và bộ máy dựng giao diện emotion-pages.js. "
        "Backend gồm một Express server, cấu hình MongoDB, hai model Mongoose, bốn nhóm route và dịch vụ gửi thư SMTP. "
        "Dữ liệu người dùng, bài chia sẻ, phản hồi và lượt thích được lưu trên MongoDB; token, email tài khoản và một số "
        "kết quả thực hành cá nhân được lưu trong LocalStorage.",
    )
    add_body(
        document,
        "Phạm vi hiện tại chưa có trang quản trị, phân quyền vai trò, chức năng xóa hoặc báo cáo bài viết, bài test "
        "cảm xúc hoàn chỉnh và bộ kiểm thử tự động. Chức năng liên hệ đã có API gửi email nhưng chỉ hoạt động khi "
        "máy chủ cấu hình SMTP; nội dung biểu mẫu không được lưu thành collection riêng.",
    )

    add_heading(document, "1.6. Các chức năng chính", 2)
    add_table(
        document,
        "Bảng 1.1. Tổng hợp chức năng chính của StellarMind",
        ["Nhóm chức năng", "Mô tả", "Trạng thái"],
        [
            ("Khám phá Hệ Mặt Trời", "Hiển thị quỹ đạo, hành tinh, tooltip, bảng thông tin, tốc độ và phím tắt.", "Đã triển khai"),
            ("Trang hành tinh", "Tám trang dùng giao diện chung, dữ liệu thiên văn, nội dung cảm xúc có nguồn và điều hướng theo mục.", "Đã triển khai"),
            ("Bài tập và tương tác", "Mỗi hành tinh có một bài thực hành riêng; một số kết quả được lưu cục bộ.", "Đã triển khai ở frontend"),
            ("Xác thực", "Đăng ký bằng username, email, mật khẩu; đăng nhập, liên kết email, bcryptjs và JWT.", "Đã có frontend và API"),
            ("Cộng đồng", "Đăng bài công khai/ẩn danh, trạm tín hiệu, trả lời, thích bài và lọc theo hành tinh.", "Đã có frontend và API"),
            ("Chatbot", "Gửi câu hỏi đến Gemini API, có phản hồi dự phòng khi thiếu khóa.", "Phụ thuộc cấu hình API"),
            ("Liên hệ", "Xác thực tài khoản, liên kết email và gửi nội dung qua dịch vụ SMTP.", "Có mã nguồn; phụ thuộc cấu hình SMTP"),
            ("Bài test", "Mục điều hướng đã xuất hiện nhưng chưa có quy trình xử lý hoàn chỉnh.", "Chưa hoàn thiện"),
        ],
        [1.45, 3.45, 1.3],
    )

    add_heading(document, "1.7. Công nghệ sử dụng", 2)
    add_body(
        document,
        "Các công nghệ dưới đây được xác định từ package, file HTML, CSS và JavaScript của dự án. Phần giải "
        "thích tập trung vào vai trò thực tế trong StellarMind thay vì mô tả chung không liên quan.",
    )
    add_table(
        document,
        "Bảng 1.2. Công nghệ và vai trò trong dự án",
        ["Công nghệ", "Vai trò trong StellarMind"],
        [
            ("HTML5", "Tổ chức cấu trúc trang chủ, trang hành tinh, cộng đồng và liên hệ."),
            ("CSS3", "Xây dựng theme không gian, responsive, glassmorphism, animation và component."),
            ("Vanilla JavaScript", "Điều khiển DOM, Canvas, trạng thái giao diện, gọi API và LocalStorage."),
            ("Canvas API", "Vẽ mô phỏng quỹ đạo, hành tinh, sao, hiệu ứng nền và hit-testing [1]."),
            ("Intersection Observer API", "Kích hoạt hiệu ứng xuất hiện theo vùng nhìn trên các trang cảm xúc [8]."),
            ("Node.js", "Môi trường thực thi JavaScript cho backend [2]."),
            ("Express.js", "Khai báo middleware, server và REST API [3]."),
            ("MongoDB / Mongoose", "Lưu User, Comment và ánh xạ schema trong Node.js [4], [5]."),
            ("JWT / bcryptjs", "Phát hành token và băm/so sánh mật khẩu [6], [7]."),
            ("Gemini API", "Sinh phản hồi chatbot từ system prompt và tin nhắn người dùng [9]."),
            ("Vercel", "Cấu hình redirect và triển khai frontend tĩnh [10]."),
            ("Nodemailer", "Tạo SMTP transporter và gửi nội dung liên hệ đến địa chỉ nhận cấu hình [11]."),
            ("Tailwind CSS", "Chỉ được nạp bằng CDN và dùng cục bộ trong trang contact.html."),
        ],
        [1.65, 4.55],
    )

    add_heading(document, "1.8. Tổng kết chương", 2)
    add_body(
        document,
        "Chương 1 đã xác định mục tiêu, phạm vi, người dùng và các module có thật trong mã nguồn StellarMind. "
        "Hệ thống có nền tảng frontend tương tác, backend Express và dữ liệu MongoDB, đồng thời vẫn còn một "
        "số chức năng ở mức giao diện hoặc phụ thuộc cấu hình triển khai.",
    )


def add_chapter_two(document):
    add_chapter(document, "CHƯƠNG 2: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")

    add_heading(document, "2.1. Khảo sát và phân tích yêu cầu", 2)
    add_body(
        document,
        "Yêu cầu được xác định bằng cách đối chiếu giao diện HTML, logic JavaScript, route Express và schema "
        "Mongoose. Cách tiếp cận này giúp phân biệt chức năng đã có luồng xử lý hoàn chỉnh với chức năng chỉ "
        "xuất hiện dưới dạng nút hoặc biểu mẫu.",
    )
    add_body(
        document,
        "Hệ thống được chia thành ba nhóm nghiệp vụ: khám phá nội dung, tương tác hỗ trợ cảm xúc và giao tiếp "
        "cộng đồng. Chatbot là dịch vụ tích hợp bên ngoài, trong khi MongoDB là thành phần lưu trữ lâu dài cho "
        "tài khoản và bình luận.",
    )

    add_heading(document, "2.2. Đối tượng sử dụng hệ thống", 2)
    add_table(
        document,
        "Bảng 2.1. Các tác nhân của hệ thống",
        ["Tác nhân", "Quyền và trách nhiệm"],
        [
            ("Khách truy cập", "Khám phá mô phỏng, xem hành tinh, sử dụng bài tập frontend, xem cộng đồng, gửi tín hiệu ẩn danh và dùng chatbot."),
            ("Người dùng đã đăng nhập", "Đăng bài công khai/ẩn danh, trả lời, thích bài, liên kết email và gửi biểu mẫu liên hệ."),
            ("Gemini API", "Nhận system prompt và tin nhắn, sau đó trả nội dung phản hồi cho route /api/chat."),
            ("Dịch vụ SMTP", "Nhận thư do Nodemailer gửi từ route /api/contact khi máy chủ đã cấu hình thông tin SMTP."),
            ("MongoDB", "Lưu trữ User, Comment và Reply; không phải tác nhân người dùng nhưng là hệ thống liên kết."),
        ],
        [1.65, 4.55],
    )

    add_heading(document, "2.3. Yêu cầu chức năng", 2)
    add_table(
        document,
        "Bảng 2.2. Yêu cầu chức năng",
        ["Mã", "Yêu cầu", "Mức ưu tiên"],
        [
            ("FR01", "Hệ thống hiển thị mô phỏng Hệ Mặt Trời và tám hành tinh.", "Cao"),
            ("FR02", "Người dùng có thể chọn hành tinh và xem bảng thông tin.", "Cao"),
            ("FR03", "Người dùng có thể mở trang chi tiết của từng hành tinh.", "Cao"),
            ("FR04", "Trang hành tinh hiển thị dấu hiệu, nguyên nhân, lời khuyên và hành động.", "Cao"),
            ("FR05", "Người dùng có thể thực hiện bài thực hành riêng của từng hành tinh.", "Trung bình"),
            ("FR06", "Người dùng có thể đăng ký bằng username, email, mật khẩu và đăng nhập.", "Cao"),
            ("FR07", "Hệ thống lưu token, username, email và cập nhật trạng thái giao diện.", "Cao"),
            ("FR08", "Người dùng đã đăng nhập có thể đăng bài công khai hoặc ẩn danh.", "Cao"),
            ("FR09", "Người dùng có thể trả lời, thích bài viết và lọc theo hành tinh.", "Cao"),
            ("FR10", "Khách có thể gửi tín hiệu ẩn danh kèm hành tinh và cường độ từ trang cảm xúc.", "Trung bình"),
            ("FR11", "Người dùng có thể gửi tin nhắn đến chatbot.", "Trung bình"),
            ("FR12", "Hệ thống trả nội dung dự phòng khi Gemini API chưa được cấu hình.", "Trung bình"),
            ("FR13", "Người dùng đã đăng nhập có thể liên kết hoặc cập nhật email tài khoản.", "Trung bình"),
            ("FR14", "Trang liên hệ kiểm tra dữ liệu và gửi email qua backend khi SMTP được cấu hình.", "Trung bình"),
        ],
        [0.65, 4.65, 0.9],
    )

    add_heading(document, "2.4. Yêu cầu phi chức năng", 2)
    add_table(
        document,
        "Bảng 2.3. Yêu cầu phi chức năng",
        ["Mã", "Nhóm", "Yêu cầu"],
        [
            ("NFR01", "Khả năng sử dụng", "Điều hướng rõ ràng, phản hồi trực quan và ngôn ngữ tiếng Việt."),
            ("NFR02", "Hiệu năng", "Canvas duy trì chuyển động mượt trên thiết bị phổ biến; ảnh có kích thước hợp lý."),
            ("NFR03", "Tương thích", "Hoạt động trên trình duyệt hiện đại hỗ trợ Canvas, Fetch và ES6."),
            ("NFR04", "Responsive", "Bố cục thích ứng với desktop, tablet và điện thoại."),
            ("NFR05", "Bảo mật", "Mật khẩu được băm; token được xác minh; bí mật phải đặt trong biến môi trường."),
            ("NFR06", "Bảo trì", "Tách CSS, JavaScript, route, model và tài nguyên theo chức năng."),
            ("NFR07", "Mở rộng", "Có thể bổ sung hành tinh, hoạt động hoặc route mà không thay đổi toàn hệ thống."),
        ],
        [0.7, 1.2, 4.3],
    )

    add_heading(document, "2.5. Biểu đồ Use Case tổng quát", 2)
    add_body(
        document,
        "Biểu đồ Use Case tổng quát mô tả phạm vi tương tác của khách truy cập, người dùng đã đăng nhập và "
        "dịch vụ Gemini API. Hệ thống hiện không có tác nhân quản trị viên vì mã nguồn chưa triển khai module quản trị.",
    )
    add_figure(document, ASSETS / "07-use-case.png", "Hình 2.1. Biểu đồ Use Case tổng quát của StellarMind")
    add_body(
        document,
        "Các chức năng khám phá và hỗ trợ cảm xúc có thể sử dụng mà không cần tài khoản. Trạm tín hiệu trên trang "
        "hành tinh cho phép gửi ẩn danh, còn bộ soạn bài chính, trả lời, thích bài và liên hệ yêu cầu đăng nhập. "
        "Chatbot kết nối đến Gemini thông qua backend để tránh đặt khóa API trực tiếp trong trình duyệt.",
    )

    add_heading(document, "2.6. Mô tả các Use Case quan trọng", 2)
    add_table(
        document,
        "Bảng 2.4. Mô tả tóm tắt các Use Case",
        ["Use Case", "Tác nhân", "Luồng chính", "Kết quả"],
        [
            ("Khám phá Hệ Mặt Trời", "Khách", "Mở trang chủ, chọn hành tinh, thay đổi tốc độ hoặc góc nhìn.", "Mô phỏng và bảng thông tin được cập nhật."),
            ("Xem trang hành tinh", "Khách", "Chọn nút khám phá hoặc liên kết hành tinh.", "Trang chủ đề hành tinh được mở."),
            ("Thực hiện bài thực hành", "Khách", "Mở mục thực hành và thao tác theo công cụ riêng của hành tinh.", "Nhận phản hồi trực quan; một số kết quả được lưu cục bộ."),
            ("Đăng ký", "Khách", "Nhập username, email, mật khẩu và gửi POST /api/auth/register.", "Tạo User và nhận JWT nếu hợp lệ."),
            ("Đăng nhập", "Khách", "Gửi thông tin đến /api/auth/login.", "Lưu token và username trong LocalStorage."),
            ("Đăng bài", "Thành viên", "Chọn chế độ, nhập nội dung và gửi bình luận.", "Comment mới được lưu."),
            ("Trả lời bài viết", "Thành viên", "Nhập phản hồi tại bài viết.", "Reply được thêm vào Comment."),
            ("Thích và lọc bài viết", "Thành viên/Khách", "Thành viên bật/tắt lượt thích; khách chọn hành tinh trên bộ lọc.", "Số lượt thích và feed được cập nhật."),
            ("Gửi tín hiệu cảm xúc", "Khách", "Nhập nội dung ẩn danh, chọn cường độ và gửi từ trang hành tinh.", "Comment gắn planet và signalLevel được lưu nếu backend sẵn sàng."),
            ("Trò chuyện chatbot", "Khách", "Nhập câu hỏi và gửi đến /api/chat.", "Hiển thị phản hồi Gemini hoặc fallback."),
            ("Gửi liên hệ", "Thành viên", "Xác thực token, liên kết email nếu cần và gửi tên, chủ đề, nội dung.", "Nodemailer gửi thư nếu SMTP được cấu hình."),
        ],
        [1.45, 0.9, 2.75, 1.1],
    )

    add_heading(document, "2.7. Kiến trúc tổng thể", 2)
    add_body(
        document,
        "StellarMind sử dụng kiến trúc tách frontend và backend. Frontend là tập hợp tài nguyên tĩnh chạy trên "
        "trình duyệt. Backend Express tiếp nhận JSON, xác thực token, thao tác MongoDB, gọi Gemini API và gửi "
        "thư liên hệ qua SMTP bằng Nodemailer.",
    )
    add_figure(document, ASSETS / "08-kien-truc.png", "Hình 2.2. Kiến trúc tổng thể của hệ thống")
    add_body(
        document,
        "Cấu trúc này cho phép triển khai frontend trên Vercel và backend trên một dịch vụ Node.js độc lập. "
        "Các module cộng đồng, trang hành tinh và liên hệ đã dùng API base có tiền tố /api; riêng chatbot.js "
        "vẫn trỏ tới domain production ở đường dẫn gốc nên cần sửa thành /api/chat trước khi triển khai hoàn chỉnh.",
    )

    add_heading(document, "2.8. Thiết kế cơ sở dữ liệu", 2)
    add_body(
        document,
        "Cơ sở dữ liệu sử dụng hai collection chính. Reply không được tách thành collection riêng mà được định "
        "nghĩa như subdocument trong mảng replies của Comment. Cách thiết kế này phù hợp với phạm vi nhỏ và "
        "giúp tải một bài viết cùng phản hồi trong một truy vấn.",
    )
    add_table(
        document,
        "Bảng 2.5. Thiết kế collection User",
        ["Trường", "Kiểu", "Ràng buộc / mục đích"],
        [
            ("username", "String", "Bắt buộc, duy nhất, viết thường, từ 3 đến 40 ký tự."),
            ("email", "String", "Bắt buộc với đăng ký mới, duy nhất, sparse, viết thường và kiểm tra định dạng."),
            ("password", "String", "Bắt buộc, tối thiểu 6 ký tự; được băm trước khi lưu."),
            ("createdAt", "Date", "Thời điểm tạo tài khoản, mặc định Date.now."),
        ],
        [1.35, 1.15, 3.7],
    )
    add_table(
        document,
        "Bảng 2.6. Thiết kế collection Comment",
        ["Trường", "Kiểu", "Ràng buộc / mục đích"],
        [
            ("author", "String", "Tên hiển thị của người đăng."),
            ("authorId", "ObjectId", "Tham chiếu User, không bắt buộc."),
            ("avatar", "String", "Ký tự đại diện."),
            ("text", "String", "Nội dung bài viết, bắt buộc."),
            ("isAnon", "Boolean", "Đánh dấu chế độ ẩn danh."),
            ("planet", "String", "Mã hành tinh liên quan, có thể null."),
            ("signalLevel", "Number", "Cường độ từ 1 đến 5, có thể null."),
            ("likedBy", "ObjectId[]", "Danh sách User đã thích bài; trường bị ẩn khỏi truy vấn mặc định."),
            ("timestamp", "Date", "Thời điểm tạo."),
            ("replies", "Reply[]", "Danh sách phản hồi lồng."),
        ],
        [1.35, 1.15, 3.7],
    )
    add_table(
        document,
        "Bảng 2.7. Thiết kế subdocument Reply",
        ["Trường", "Kiểu", "Mục đích"],
        [
            ("author", "String", "Tên người phản hồi."),
            ("authorId", "ObjectId", "Tham chiếu User nếu đã đăng nhập."),
            ("avatar", "String", "Ký tự đại diện."),
            ("text", "String", "Nội dung phản hồi."),
            ("timestamp", "Date", "Thời điểm phản hồi."),
        ],
        [1.35, 1.15, 3.7],
    )

    add_heading(document, "2.9. Mô hình dữ liệu", 2)
    add_figure(document, ASSETS / "09-mo-hinh-du-lieu.png", "Hình 2.3. Mô hình dữ liệu User, Comment và Reply")
    add_body(
        document,
        "Quan hệ giữa User và Comment được thể hiện qua authorId và mảng likedBy. Khi tín hiệu được gửi ẩn danh từ "
        "trang hành tinh, authorId có thể bằng null; ngược lại, phản hồi và lượt thích yêu cầu User hợp lệ. Thiết kế "
        "này hỗ trợ trải nghiệm ẩn danh nhưng vẫn cần cơ chế kiểm duyệt và giới hạn tần suất ở phiên bản sau.",
    )

    add_heading(document, "2.10. Biểu đồ tuần tự", 2)
    add_heading(document, "2.10.1. Luồng đăng ký và đăng nhập", 3)
    add_body(
        document,
        "Giao diện cộng đồng gửi username, email và password khi đăng ký; đăng nhập dùng username và password. "
        "Backend truy vấn User, sử dụng bcryptjs để băm hoặc so sánh mật khẩu, sau đó tạo JWT có thời hạn 30 ngày.",
    )
    add_figure(document, ASSETS / "10-tuan-tu-xac-thuc.png", "Hình 2.4. Biểu đồ tuần tự đăng ký và đăng nhập")

    add_heading(document, "2.10.2. Luồng đăng bài và trả lời", 3)
    add_body(
        document,
        "Frontend gửi Bearer token khi người dùng đã đăng nhập. Tạo Comment cho phép token tùy chọn để hỗ trợ trạm "
        "tín hiệu ẩn danh; route reply và like bắt buộc xác thực. Backend chuẩn hóa dữ liệu trước khi lưu Comment, "
        "Reply hoặc cập nhật mảng likedBy trong MongoDB.",
    )
    add_figure(document, ASSETS / "11-tuan-tu-cong-dong.png", "Hình 2.5. Biểu đồ tuần tự đăng bài và trả lời")

    add_heading(document, "2.10.3. Luồng chatbot", 3)
    add_body(
        document,
        "Chatbot UI chỉ gửi nội dung tin nhắn. System prompt và khóa Gemini được xử lý tại backend. Route chat "
        "thử nhiều model theo thứ tự dự phòng và trả nội dung đầu tiên hợp lệ.",
    )
    add_figure(document, ASSETS / "12-tuan-tu-chatbot.png", "Hình 2.6. Biểu đồ tuần tự tương tác chatbot")

    add_heading(document, "2.10.4. Luồng gửi liên hệ", 3)
    add_body(
        document,
        "Người dùng phải có JWT hợp lệ và email liên kết. Frontend có thể gọi PATCH /api/auth/email trước khi gửi "
        "POST /api/contact. Backend kiểm tra giới hạn độ dài, áp dụng thời gian chờ 60 giây theo người dùng và "
        "gọi Nodemailer để chuyển thư đến dịch vụ SMTP.",
    )
    add_figure(document, ASSETS / "14-tuan-tu-lien-he.png", "Hình 2.7. Biểu đồ tuần tự gửi biểu mẫu liên hệ")

    add_heading(document, "2.11. Thiết kế giao diện người dùng", 2)
    add_heading(document, "2.11.1. Trang chủ mô phỏng", 3)
    add_body(
        document,
        "Trang chủ ưu tiên vùng Canvas toàn màn hình, thanh điều hướng, danh sách tám hành tinh và bảng điều "
        "khiển tốc độ. Các lớp nền, quỹ đạo và ánh sáng tạo cảm giác không gian nhưng vẫn giữ thông tin chính "
        "ở vùng dễ quan sát.",
    )
    add_figure(document, ASSETS / "01-trang-chu.png", "Hình 2.8. Giao diện trang chủ mô phỏng Hệ Mặt Trời")

    add_heading(document, "2.11.2. Trang hành tinh", 3)
    add_body(
        document,
        "Trang hành tinh sử dụng màu chủ đề, hình texture lớn, thông tin thiên văn và menu neo theo chiều cuộn. "
        "Các khối dùng chung gồm tổng quan, dấu hiệu, bản đồ cơ thể, vòng lặp cảm xúc, nguồn tham khảo, tình huống "
        "thực tế, thực hành, trạm tín hiệu và bài chia sẻ gần đây.",
    )
    add_figure(document, ASSETS / "03-trang-hanh-tinh.png", "Hình 2.9. Giao diện trang hành tinh Trái Đất")

    add_heading(document, "2.11.3. Bài thực hành và trạm tín hiệu", 3)
    add_body(
        document,
        "Mỗi hành tinh cung cấp một công cụ thực hành phù hợp với chủ đề thay vì dùng cùng một bài tập cho mọi trang. "
        "Trạm tín hiệu cho phép người dùng ghi nội dung ẩn danh, chọn cường độ từ 1 đến 5 và chuyển tiếp sang cộng đồng "
        "với bộ lọc hành tinh tương ứng.",
    )
    add_figure(document, ASSETS / "03b-thuc-hanh.png", "Hình 2.10. Giao diện bài thực hành tiếp đất trên trang Trái Đất")
    add_figure(document, ASSETS / "03c-tram-tin-hieu.png", "Hình 2.11. Giao diện trạm tín hiệu cảm xúc")

    add_heading(document, "2.11.4. Trang cộng đồng và xác thực", 3)
    add_body(
        document,
        "Trang cộng đồng có vùng soạn bài, hai chế độ hiển thị danh tính, bộ lọc theo hành tinh, nhãn cường độ tín hiệu, "
        "nút thích và vùng trả lời. Giao diện xác thực được trình bày dưới dạng modal; tab đăng ký yêu cầu thêm email.",
    )
    add_figure(document, ASSETS / "04-cong-dong.png", "Hình 2.12. Giao diện trang cộng đồng")
    add_figure(document, ASSETS / "05-dang-nhap.png", "Hình 2.13. Giao diện đăng nhập và đăng ký có email")

    add_heading(document, "2.11.5. Cửa sổ chatbot", 3)
    add_body(
        document,
        "Chatbot được chèn như một module dùng chung trên nhiều trang. Cửa sổ có lời chào, gợi ý câu hỏi, vùng "
        "tin nhắn và trạng thái đang nhập. Thiết kế tách biệt này giúp tái sử dụng mà không thay đổi cấu trúc trang.",
    )
    add_figure(document, ASSETS / "02-chatbot.png", "Hình 2.14. Cửa sổ Trợ lý Tinh hà trên trang chủ")

    add_heading(document, "2.11.6. Trang liên hệ", 3)
    add_body(
        document,
        "Trang liên hệ giới thiệu đội ngũ, thông tin liên lạc và biểu mẫu. Frontend yêu cầu đăng nhập, kiểm tra dữ liệu, "
        "đồng bộ email tài khoản khi cần và gọi API gửi thư. Backend không tạo collection Contact; khả năng gửi thực tế "
        "phụ thuộc cấu hình SMTP trên máy chủ.",
    )
    add_figure(document, ASSETS / "06-lien-he.png", "Hình 2.15. Giao diện trang liên hệ")

    add_heading(document, "2.12. Tổng kết chương", 2)
    add_body(
        document,
        "Chương 2 đã mô hình hóa các tác nhân, yêu cầu, Use Case, kiến trúc, dữ liệu và luồng xử lý cốt lõi. "
        "Thiết kế cho thấy frontend, backend, MongoDB và Gemini được phân tách rõ, đồng thời phản ánh đúng các "
        "giới hạn hiện tại như thiếu quản trị, cấu hình SMTP và một số đường dẫn production chưa hoàn chỉnh.",
    )


def add_chapter_three(document):
    add_chapter(document, "CHƯƠNG 3: XÂY DỰNG, TRIỂN KHAI VÀ KIỂM THỬ")

    add_heading(document, "3.1. Môi trường phát triển", 2)
    add_table(
        document,
        "Bảng 3.1. Môi trường và công cụ phát triển",
        ["Thành phần", "Sử dụng"],
        [
            ("Visual Studio Code", "Soạn thảo HTML, CSS, JavaScript và cấu hình dự án."),
            ("Git", "Quản lý phiên bản mã nguồn [12]."),
            ("Trình duyệt hiện đại", "Chạy Canvas, kiểm tra DOM, Fetch API và responsive."),
            ("Node.js / npm", "Cài dependency và chạy Express server."),
            ("MongoDB", "Lưu dữ liệu cục bộ hoặc từ dịch vụ MongoDB tương thích."),
            ("Live Server / HTTP server", "Phục vụ frontend qua HTTP để tránh hạn chế khi mở file trực tiếp."),
        ],
        [1.75, 4.45],
    )

    add_heading(document, "3.2. Cấu trúc thư mục dự án", 2)
    add_body(
        document,
        "Mã nguồn được chia thành frontend và backend. Frontend tiếp tục chia theo HTML, CSS, JavaScript, ảnh "
        "minh họa và texture. Backend chia thành config, models, routes và services; thư mục services hiện chứa "
        "module cấu hình và gửi thư bằng Nodemailer.",
    )
    add_figure(document, ASSETS / "13-cau-truc-thu-muc.png", "Hình 3.1. Cấu trúc thư mục chính của dự án")
    add_body(
        document,
        "Trong workspace hiện còn một số bản sao JavaScript trong frontend/html/js và frontend/js có nội dung "
        "khác nhau. Báo cáo lấy các file được trang HTML chính tham chiếu trong frontend/js làm căn cứ. Việc "
        "xóa hoặc đồng bộ các bản sao là nhiệm vụ bảo trì cần thực hiện sau đồ án.",
    )

    add_heading(document, "3.3. Xây dựng mô phỏng Hệ Mặt Trời bằng Canvas", 2)
    for text in [
        (
            "Module app.js định nghĩa mảng PLANETS chứa bán kính hiển thị, bán kính quỹ đạo, tốc độ, màu sắc, "
            "góc ban đầu, thông số thiên văn và mô tả. Trạng thái ứng dụng theo dõi hành tinh đang chọn, tốc độ, "
            "camera, con trỏ, parallax và thời gian mô phỏng."
        ),
        (
            "Canvas nền vẽ các lớp thiên hà, sao tĩnh và sao parallax. Canvas mô phỏng vẽ Mặt Trời, quỹ đạo, hành "
            "tinh, vành đai tiểu hành tinh và sao băng. Vòng lặp requestAnimationFrame cập nhật góc quay và vị trí "
            "theo hệ số tốc độ."
        ),
        (
            "Hệ thống chuyển đổi tọa độ giữa không gian mô phỏng và màn hình để thực hiện hit-testing. Khi người "
            "dùng chọn một hành tinh, camera nội suy đến vị trí mục tiêu, bảng thông tin xuất hiện và các hành tinh "
            "không liên quan được làm mờ."
        ),
    ]:
        add_body(document, text)

    add_heading(document, "3.4. Xây dựng trang nội dung cho tám hành tinh", 2)
    add_body(
        document,
        "Tám trang hành tinh được rút gọn thành các lớp vỏ HTML khai báo thuộc tính data-planet và nạp ba tài nguyên "
        "dùng chung: emotion-pages.css, planet-data.js và emotion-pages.js. planet-data.js lưu dữ liệu thiên văn, "
        "ánh xạ cảm xúc, dấu hiệu, nghiên cứu, tình huống và cấu hình bài thực hành; emotion-pages.js dựng giao diện "
        "và gắn sự kiện theo dữ liệu của hành tinh hiện tại.",
    )
    add_body(
        document,
        "Các hành tinh được gắn với chủ đề ở frontend như sau: Sao Thủy – cô đơn, Sao Kim – yêu thương, Trái Đất – "
        "lo âu, Sao Hỏa – giận dữ, Sao Mộc – áp lực, Sao Thổ – trầm cảm, Sao Thiên Vương – tự ti và Sao Hải Vương – "
        "mất phương hướng. Đây là hệ ánh xạ cần được đồng bộ lại với system prompt của backend.",
    )

    add_heading(document, "3.5. Bài thực hành và hoạt động tương tác", 2)
    add_body(
        document,
        "Phiên bản mới không dùng một bài tập chung cho mọi hành tinh. Bộ máy emotion-pages.js chọn một trong tám "
        "loại thực hành theo cấu hình practice.type, tạo giao diện và xử lý trạng thái hoàn toàn bằng Vanilla JavaScript.",
    )
    add_body(
        document,
        "Các thực hành gồm: thang kết nối ba nấc ở Sao Thủy; quỹ đạo yêu thương có ranh giới ở Sao Kim; tiếp đất "
        "5–4–3–2–1 kết hợp nhịp thở 4–6 ở Trái Đất; buồng hạ nhiệt 90 giây ở Sao Hỏa; ba quỹ đạo ưu tiên ở Sao Mộc; "
        "vi hành động theo mức năng lượng ở Sao Thổ; tái cấu trúc suy nghĩ và ghi nhận thành tựu ở Sao Thiên Vương; "
        "la bàn ba giá trị ở Sao Hải Vương. Một số kết quả được lưu trong LocalStorage để người dùng xem lại.",
    )

    add_heading(document, "3.6. Xây dựng hệ thống đăng ký và đăng nhập", 2)
    add_body(
        document,
        "Route đăng ký yêu cầu username, email và password; kiểm tra định dạng email, username hoặc email trùng trước "
        "khi tạo User. Middleware pre-save trong model sử dụng bcryptjs để tạo salt và băm mật khẩu. Route đăng nhập "
        "gọi matchPassword trước khi phát hành JWT; PATCH /api/auth/email cho phép tài khoản cũ liên kết email.",
    )
    add_body(
        document,
        "Frontend lưu token, username và email trong LocalStorage với các khóa stellar_auth_token, "
        "stellar_auth_username và stellar_auth_email. Các trang cộng đồng và liên hệ đọc các giá trị này để hiển thị "
        "trạng thái đăng nhập, gửi Bearer token và đồng bộ email tài khoản khi cần.",
    )

    add_heading(document, "3.7. Xây dựng cộng đồng chia sẻ cảm xúc", 2)
    add_body(
        document,
        "Trang cộng đồng tải danh sách Comment theo thời gian giảm dần. Mỗi bài hiển thị tác giả, avatar, thời gian, "
        "hành tinh, mức tín hiệu, nội dung, số phản hồi, số lượt thích và trạng thái likedByMe. Nội dung do người dùng "
        "nhập được escape trước khi chèn vào HTML để hạn chế chèn mã trực tiếp.",
    )
    add_body(
        document,
        "Ở bộ soạn bài chính, giao diện yêu cầu đăng nhập. Trạm tín hiệu trên trang hành tinh lại cho phép gửi ẩn danh "
        "vì POST /api/comments chấp nhận token tùy chọn; reply và like bắt buộc JWT. GET /api/comments hỗ trợ query "
        "planet, nhưng trang cộng đồng hiện vẫn tải toàn bộ danh sách rồi lọc phía frontend.",
    )

    add_heading(document, "3.8. Tích hợp chatbot Gemini", 2)
    add_body(
        document,
        "Module chatbot.js tạo giao diện dùng chung và gửi message đến /api/chat. Backend ghép system prompt, một "
        "phản hồi xác nhận vai trò và tin nhắn người dùng thành mảng contents của Gemini API.",
    )
    add_body(
        document,
        "Route thử lần lượt một danh sách model để xử lý trường hợp model trả về 404. Nếu không có GEMINI_API_KEY, "
        "backend trả một phản hồi hướng dẫn cục bộ. Cơ chế này giúp giao diện không bị treo nhưng không thay thế "
        "việc cấu hình và kiểm thử model hợp lệ trước khi triển khai.",
    )
    add_body(
        document,
        "System prompt hiện còn dùng ánh xạ cảm xúc cũ và một số thuật ngữ thiên về điều trị, chưa đồng bộ với dữ liệu "
        "mới của tám trang hành tinh. Ngoài ra, URL production trong chatbot.js đang thiếu /api/chat; vì vậy đây là "
        "hai điểm phải hiệu chỉnh trước khi đánh giá chatbot trong môi trường production.",
    )
    add_callout(
        document,
        "Nguyên tắc sử dụng",
        "Phản hồi của chatbot phải được trình bày là gợi ý hỗ trợ cảm xúc, không phải tư vấn y khoa. Với tình huống "
        "khẩn cấp hoặc nguy cơ gây hại, sản phẩm tương lai cần hiển thị hướng dẫn tìm trợ giúp chuyên môn.",
    )

    add_heading(document, "3.9. Xây dựng và sử dụng REST API", 2)
    add_table(
        document,
        "Bảng 3.2. Danh sách REST API",
        ["Phương thức", "Endpoint", "Chức năng", "Xác thực"],
        [
            ("POST", "/api/auth/register", "Tạo tài khoản bằng username, email, password và trả JWT.", "Không"),
            ("POST", "/api/auth/login", "Kiểm tra tài khoản và trả JWT.", "Không"),
            ("PATCH", "/api/auth/email", "Liên kết hoặc cập nhật email tài khoản.", "Bắt buộc"),
            ("GET", "/api/comments", "Lấy bình luận; hỗ trợ query planet và trạng thái lượt thích.", "Token tùy chọn"),
            ("POST", "/api/comments", "Tạo bình luận mới.", "Token tùy chọn ở backend"),
            ("POST", "/api/comments/:id/reply", "Thêm phản hồi vào bình luận.", "Bắt buộc"),
            ("POST", "/api/comments/:id/like", "Bật hoặc tắt lượt thích của người dùng.", "Bắt buộc"),
            ("POST", "/api/chat", "Gửi tin nhắn tới Gemini hoặc fallback.", "Không"),
            ("GET", "/api/contact/status", "Kiểm tra máy chủ đã cấu hình SMTP hay chưa.", "Không"),
            ("POST", "/api/contact", "Gửi biểu mẫu liên hệ qua Nodemailer và SMTP.", "Bắt buộc"),
        ],
        [0.8, 1.75, 2.8, 0.85],
    )
    add_body(
        document,
        "Có sự khác biệt có chủ đích giữa hai luồng tạo nội dung: bộ soạn bài cộng đồng yêu cầu đăng nhập, còn trạm "
        "tín hiệu cho phép gửi ẩn danh. Do cùng dùng POST /api/comments, backend vẫn chấp nhận request không có token. "
        "Khi triển khai thực tế cần giới hạn tần suất và kiểm duyệt để tránh lạm dụng luồng ẩn danh.",
    )

    add_heading(document, "3.10. Triển khai hệ thống", 2)
    add_body(
        document,
        "Frontend có file vercel.json chuyển hướng đường dẫn gốc sang /html/index.html. Backend đọc MONGO_URI, "
        "JWT_SECRET, GEMINI_API_KEY và nhóm biến SMTP từ môi trường; nếu không có MONGO_URI, hệ thống thử MongoDB cục bộ.",
    )
    add_body(
        document,
        "Cấu hình production trong community.js, emotion-pages.js và trang liên hệ đã dùng domain Render kèm /api. "
        "Riêng chatbot.js vẫn trỏ tới đường dẫn gốc. Bước triển khai chính thức phải sửa URL này, cấu hình SMTP, giới "
        "hạn CORS theo domain frontend và loại bỏ JWT secret dự phòng được viết trực tiếp trong mã nguồn.",
    )

    add_heading(document, "3.11. Kiểm thử hệ thống", 2)
    add_body(
        document,
        "Trong phạm vi cập nhật báo cáo, 18 file JavaScript đã được kiểm tra cú pháp bằng node --check. Trang chủ, "
        "trang Trái Đất, bài thực hành, trạm tín hiệu, cộng đồng, modal đăng ký có email và trang liên hệ đã được mở "
        "trên máy chủ HTTP cục bộ. Các luồng phụ thuộc MongoDB, SMTP hoặc Gemini chưa được xác nhận đầu cuối.",
    )
    add_table(
        document,
        "Bảng 3.3. Kịch bản kiểm thử và trạng thái xác minh",
        ["Mã", "Kịch bản", "Kết quả mong đợi", "Trạng thái"],
        [
            ("TC01", "Mở trang chủ", "Canvas và danh sách tám hành tinh hiển thị.", "Đã quan sát"),
            ("TC02", "Mở trang Trái Đất", "Hero, menu mục, nội dung cảm xúc và nguồn tham khảo hiển thị.", "Đã quan sát"),
            ("TC03", "Mở bài thực hành và trạm tín hiệu", "Công cụ tiếp đất và form cường độ tín hiệu hiển thị.", "Đã quan sát"),
            ("TC04", "Mở trang cộng đồng", "Composer, bộ lọc, nhãn tín hiệu và nút xác thực hiển thị.", "Đã quan sát"),
            ("TC05", "Mở modal đăng ký", "Form username, email và password xuất hiện.", "Đã quan sát"),
            ("TC06", "Kiểm tra cú pháp 18 file JavaScript", "Không có SyntaxError.", "Đạt kiểm tra cú pháp"),
            ("TC07", "Đăng ký và đăng nhập", "Tạo User, nhận token, lưu username và email.", "Cần kiểm tra với MongoDB"),
            ("TC08", "Đăng bài, gửi tín hiệu và trả lời", "Comment/Reply được lưu và tải lại.", "Cần kiểm tra với MongoDB"),
            ("TC09", "Bật/tắt lượt thích", "likeCount và likedByMe cập nhật đúng.", "Cần kiểm tra với MongoDB"),
            ("TC10", "Gửi chatbot", "Nhận phản hồi Gemini qua /api/chat.", "Cần sửa URL production và có API key"),
            ("TC11", "Gửi liên hệ", "Email được chuyển qua SMTP và áp dụng thời gian chờ.", "Cần cấu hình SMTP"),
            ("TC12", "Mất kết nối backend", "Giao diện hiển thị thông báo lỗi phù hợp.", "Cần kiểm tra"),
            ("TC13", "Responsive", "Không tràn hoặc che khuất nội dung.", "Cần kiểm tra đa thiết bị"),
            ("TC14", "Triển khai production", "Mọi module gọi đúng endpoint /api.", "Chatbot cần hoàn thiện"),
        ],
        [0.55, 2.1, 2.4, 1.15],
    )

    add_heading(document, "3.12. Kết quả đạt được", 2)
    add_table(
        document,
        "Bảng 3.4. Mức độ hoàn thành các nhóm chức năng",
        ["Nhóm", "Kết quả"],
        [
            ("Mô phỏng và điều hướng", "Có thể chạy cục bộ; giao diện và Canvas hiển thị đúng."),
            ("Tám trang hành tinh", "Đã dùng chung bộ máy dựng trang, dữ liệu theo hành tinh và tám bài thực hành riêng."),
            ("Cộng đồng và xác thực", "Đã bổ sung email, trạm tín hiệu, lượt thích, phản hồi và lọc; cần kiểm thử đầu cuối với MongoDB."),
            ("Chatbot", "Đã có UI, route, prompt và fallback; cần xác nhận model và API key."),
            ("Liên hệ", "Đã có validation, xác thực, đồng bộ email và gửi SMTP; hoạt động thực tế phụ thuộc cấu hình máy chủ."),
            ("Kiểm thử", "Đã kiểm tra cú pháp và hiển thị chọn lọc; chưa có test tự động."),
        ],
        [2.0, 4.2],
    )
    add_body(
        document,
        "Kết quả chính của đồ án là một hệ thống web có bản sắc giao diện rõ, kết hợp Canvas, nội dung đa trang, "
        "REST API, xác thực và cơ sở dữ liệu. Mã nguồn thể hiện khả năng tích hợp nhiều kỹ thuật nhưng vẫn cần một "
        "giai đoạn ổn định hóa trước khi triển khai công khai.",
    )

    add_heading(document, "3.13. Hạn chế hiện tại", 2)
    add_bullets(
        document,
        [
            "URL production của chatbot.js chưa trỏ tới /api/chat.",
            "Nội dung ánh xạ cảm xúc giữa frontend và system prompt chatbot chưa đồng bộ.",
            "Chưa có trang quản trị, xóa bài, báo cáo nội dung hoặc kiểm duyệt.",
            "Backend dùng JWT secret dự phòng trong mã nguồn và CORS chưa giới hạn domain.",
            "Luồng bình luận ẩn danh chưa có rate limiting; thời gian chờ liên hệ chỉ lưu trong bộ nhớ và mất khi server khởi động lại.",
            "Chức năng liên hệ chỉ hoạt động khi SMTP được cấu hình và không lưu lịch sử gửi vào cơ sở dữ liệu.",
            "Mục bài test chưa được triển khai hoàn chỉnh.",
            "Một số số liệu thiên văn và nội dung cảm xúc vẫn cần tiếp tục rà soát nguồn.",
            "Có các bản sao file JavaScript không đồng nhất và chưa có kiểm thử tự động.",
        ],
    )

    add_heading(document, "3.14. Tổng kết chương", 2)
    add_body(
        document,
        "Chương 3 đã trình bày cách xây dựng từng module, danh sách API, phạm vi triển khai và trạng thái kiểm thử. "
        "Các kết quả được phân loại theo bằng chứng quan sát hoặc điều kiện cần kiểm tra, tránh khẳng định những "
        "luồng chưa được chạy đầu cuối.",
    )


def add_conclusion(document):
    add_chapter(document, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    add_heading(document, "1. Kết quả đạt được", 2)
    add_body(
        document,
        "Dự án đã xây dựng được website StellarMind với mô phỏng Hệ Mặt Trời, tám trang hành tinh dùng chung bộ máy "
        "giao diện, tám bài thực hành, trạm tín hiệu, cộng đồng có lượt thích, xác thực bằng email và chatbot. "
        "Sản phẩm vận dụng Canvas, lập trình giao diện, REST API, MongoDB, JWT, Gemini và gửi thư qua SMTP.",
    )
    add_body(
        document,
        "Quá trình thực hiện giúp củng cố kỹ năng phân tích yêu cầu, tổ chức mã nguồn, thiết kế dữ liệu, xử lý "
        "trạng thái phía trình duyệt và xây dựng giao tiếp frontend–backend. Báo cáo cũng xác định rõ giới hạn "
        "của sản phẩm để tránh mô tả vượt quá mức độ hoàn thiện thực tế.",
    )

    add_heading(document, "2. Hạn chế", 2)
    add_body(
        document,
        "Hệ thống chưa có đầy đủ điều kiện kiểm thử đầu cuối, chatbot chưa hoàn thiện URL production và còn thiếu "
        "các chức năng quản trị, kiểm duyệt và bài test. Chức năng liên hệ phụ thuộc SMTP, chưa lưu lịch sử; nội dung "
        "cảm xúc giữa frontend và chatbot cần được đồng bộ thuật ngữ và tiếp tục rà soát nguồn.",
    )

    add_heading(document, "3. Hướng phát triển", 2)
    add_bullets(
        document,
        [
            "Hoàn thiện bài test cảm xúc và giải thích rõ cơ sở của kết quả.",
            "Đồng bộ ánh xạ cảm xúc giữa frontend, chatbot và tài liệu.",
            "Xây dựng trang quản trị, cơ chế kiểm duyệt, xóa và báo cáo nội dung.",
            "Hoàn thiện cấu hình SMTP, lưu lịch sử hoặc trạng thái liên hệ khi có yêu cầu nghiệp vụ.",
            "Bổ sung Jest, Supertest hoặc Playwright cho kiểm thử tự động.",
            "Tối ưu Canvas, ảnh và animation trên thiết bị cấu hình thấp.",
            "Đưa JWT secret, API key và CORS vào cấu hình môi trường an toàn.",
            "Bổ sung rate limiting, validation và chính sách xác thực thống nhất ở backend.",
            "Thêm nguồn chính thức cho dữ liệu thiên văn và nguồn chuyên môn cho nội dung cảm xúc.",
            "Hoàn thiện URL production, logging và giám sát lỗi khi triển khai.",
        ],
    )


def add_references(document):
    add_chapter(document, "TÀI LIỆU THAM KHẢO")
    references = [
        ("[1] MDN Web Docs. Canvas API.", "https://developer.mozilla.org/en-US/docs/Web/API/Canvas_API"),
        ("[2] OpenJS Foundation. Node.js Documentation.", "https://nodejs.org/en/docs/"),
        ("[3] Express.js. Express 4.x API Reference.", "https://expressjs.com/en/4x/api.html"),
        ("[4] MongoDB, Inc. MongoDB Documentation.", "https://www.mongodb.com/docs/"),
        ("[5] Mongoose. Mongoose Documentation.", "https://mongoosejs.com/docs/"),
        ("[6] Auth0. node-jsonwebtoken.", "https://github.com/auth0/node-jsonwebtoken"),
        ("[7] dcodeIO. bcrypt.js.", "https://github.com/dcodeIO/bcrypt.js/"),
        ("[8] MDN Web Docs. Intersection Observer API.", "https://developer.mozilla.org/en-US/docs/Web/API/Intersection_Observer_API"),
        ("[9] Google AI for Developers. Gemini API Documentation.", "https://ai.google.dev/gemini-api/docs/api-overview"),
        ("[10] Vercel. Vercel Documentation.", "https://vercel.com/docs/"),
        ("[11] Nodemailer. Nodemailer Documentation.", "https://nodemailer.com/"),
        ("[12] Git. Git Reference.", "https://git-scm.com/docs"),
        ("[13] National Institute of Mental Health. Mental Health Information.", "https://www.nimh.nih.gov/health"),
        ("[14] World Health Organization. Mental health.", "https://www.who.int/health-topics/mental-health"),
        ("[15] American Psychological Association. Topics.", "https://www.apa.org/topics"),
        ("[16] National Health Service. Mental health.", "https://www.nhs.uk/mental-health/"),
    ]
    for label, url in references:
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.hanging_indent = Cm(0.5)
        run = p.add_run(f"{label} Truy cập ngày 12/06/2026, từ {url}")
        set_run_font(run, size=12)


def add_appendices(document):
    add_chapter(document, "PHỤ LỤC")
    add_heading(document, "Phụ lục A. Cấu trúc thư mục rút gọn", 2)
    code = """Project_Web_BWD/
├── frontend/
│   ├── html/
│   ├── css/
│   ├── js/
│   ├── assets/images/planets/
│   ├── planets/
│   └── vercel.json
└── backend/
    ├── config/db.js
    ├── models/User.js
    ├── models/Comment.js
    ├── routes/auth.js
    ├── routes/comments.js
    ├── routes/chat.js
    ├── routes/contact.js
    ├── services/mailer.js
    ├── server.js
    └── package.json"""
    p = document.add_paragraph(style="Code")
    p.add_run(code)

    add_heading(document, "Phụ lục B. Hướng dẫn cài đặt và chạy", 2)
    add_numbered(
        document,
        [
            "Cài Node.js và MongoDB hoặc chuẩn bị MONGO_URI từ dịch vụ MongoDB.",
            "Mở thư mục backend và chạy npm install.",
            "Tạo backend/.env với MONGO_URI, JWT_SECRET, GEMINI_API_KEY và cấu hình SMTP khi dùng trang liên hệ.",
            "Chạy npm start hoặc npm run dev trong thư mục backend.",
            "Phục vụ thư mục frontend bằng Live Server hoặc một HTTP server.",
            "Mở /html/index.html trên trình duyệt và kiểm tra các route API.",
        ],
    )

    add_heading(document, "Phụ lục C. Biến môi trường", 2)
    p = document.add_paragraph(style="Code")
    p.add_run(
        "MONGO_URI=mongodb://127.0.0.1:27017/stellarmind\n"
        "JWT_SECRET=<chuỗi_bí_mật_mạnh>\n"
        "GEMINI_API_KEY=<khóa_api_gemini>\n"
        "SMTP_HOST=<máy_chủ_smtp>\n"
        "SMTP_PORT=587\n"
        "SMTP_SECURE=false\n"
        "SMTP_USER=<tài_khoản_smtp>\n"
        "SMTP_PASS=<mật_khẩu_ứng_dụng>\n"
        "MAIL_FROM=<địa_chỉ_gửi>\n"
        "MAIL_TO=<địa_chỉ_nhận>\n"
        "PORT=5000"
    )

    add_heading(document, "Phụ lục D. Một số đoạn mã tiêu biểu", 2)
    add_body(
        document,
        "Đoạn mã sau thể hiện cách server gắn bốn nhóm route chính. Nội dung được rút gọn từ backend/server.js.",
    )
    p = document.add_paragraph(style="Code")
    p.add_run(
        "app.use(cors());\n"
        "app.use(express.json());\n"
        "app.use('/api/auth', require('./routes/auth'));\n"
        "app.use('/api/comments', require('./routes/comments'));\n"
        "app.use('/api/chat', require('./routes/chat'));\n"
        "app.use('/api/contact', require('./routes/contact'));"
    )

    add_body(
        document,
        "Model User dùng hook pre-save để băm mật khẩu trước khi lưu.",
    )
    p = document.add_paragraph(style="Code")
    p.add_run(
        "UserSchema.pre('save', async function(next) {\n"
        "  if (!this.isModified('password')) return next();\n"
        "  const salt = await bcrypt.genSalt(10);\n"
        "  this.password = await bcrypt.hash(this.password, salt);\n"
        "});"
    )


def build():
    document = Document()
    configure_styles(document)
    for section in document.sections:
        set_section_geometry(section)
    set_update_fields(document)

    add_cover(document, secondary=False)
    document.add_page_break()
    add_cover(document, secondary=True)
    add_front_matter(document)
    add_introduction(document)
    add_chapter_one(document)
    add_chapter_two(document)
    add_chapter_three(document)
    add_conclusion(document)
    add_references(document)
    add_appendices(document)

    core = document.core_properties
    core.title = TITLE
    core.subject = "Báo cáo dự án StellarMind"
    core.author = STUDENT
    core.keywords = "StellarMind, Canvas, Express, MongoDB, Gemini, Nodemailer, báo cáo dự án"
    core.comments = "Báo cáo được xây dựng từ mã nguồn thực tế của dự án StellarMind."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
