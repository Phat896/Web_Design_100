import json
import sys
from pathlib import Path

from docx import Document


def clean_text(text):
    return " ".join(text.replace("\xa0", " ").split())


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: inspect_docx.py INPUT.docx OUTPUT.json")

    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    document = Document(source)

    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs):
        text = clean_text(paragraph.text)
        if not text:
            continue
        paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else "",
                "text": text,
            }
        )

    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([clean_text(cell.text) for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    sections = []
    for index, section in enumerate(document.sections):
        sections.append(
            {
                "index": index,
                "width": section.page_width,
                "height": section.page_height,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
            }
        )

    payload = {
        "source": str(source),
        "paragraph_count": len(document.paragraphs),
        "nonempty_paragraph_count": len(paragraphs),
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
        "section_count": len(document.sections),
        "paragraphs": paragraphs,
        "tables": tables,
        "sections": sections,
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
